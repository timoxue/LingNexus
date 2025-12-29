"""
测试代码执行器
"""

import sys
import io
from pathlib import Path

# Windows 编码修复
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

from lingnexus.utils.code_executor import extract_and_execute_code


def test_simple_code():
    """测试简单代码执行"""
    print("=" * 60)
    print("测试简单代码执行")
    print("=" * 60)
    
    code = """
from docx import Document

doc = Document()
doc.add_heading('Test', level=1)
doc.add_paragraph('This is a test')
doc.save('test_simple.docx')
print("File created")
"""
    
    response_text = f"```python{code}```"
    
    result = extract_and_execute_code(response_text)
    
    print(f"成功: {result['success']}")
    print(f"错误: {result.get('error', 'None')}")
    print(f"输出: {result.get('output', 'None')}")
    
    # 检查文件
    test_file = Path("test_simple.docx")
    if test_file.exists():
        print(f"\n✅ 文件创建成功: {test_file}")
        print(f"   文件大小: {test_file.stat().st_size} 字节")
    else:
        print(f"\n❌ 文件未创建: {test_file}")


if __name__ == "__main__":
    test_simple_code()

