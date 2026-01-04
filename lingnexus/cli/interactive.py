"""
交互式测试工具
提供用户友好的命令行交互界面
"""

import sys
import io
import os
import asyncio
from pathlib import Path
from typing import Optional, Dict, Any

# Windows 编码修复
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

from ..config import init_agentscope, ModelType
from ..agent import create_progressive_agent  # 通过 react_agent.py 作为统一入口，使用渐进式披露
from ..utils.code_executor import extract_and_execute_code_async, extract_and_execute_multi_language
from ..utils.code_executor import extract_code_from_text
from agentscope.message import Msg


class InteractiveTester:
    """交互式测试工具"""
    
    def __init__(
        self,
        model_type: ModelType = ModelType.QWEN,
        model_name: Optional[str] = None,
        auto_execute_code: bool = True,
        enable_studio: bool = False,
    ):
        """
        初始化交互式测试工具
        
        Args:
            model_type: 模型类型
            model_name: 模型名称
            auto_execute_code: 是否自动执行代码
            enable_studio: 是否启用 Studio
        """
        self.model_type = model_type
        self.model_name = model_name or ("qwen-max" if model_type == ModelType.QWEN else "deepseek-chat")
        self.auto_execute_code = auto_execute_code
        self.enable_studio = enable_studio
        
        self.agent: Optional[Any] = None
        self.conversation_history = []
        self.current_mode = "chat"  # chat, test, help
        
    def _init_agentscope(self):
        """初始化 AgentScope"""
        # 如果启用 Studio，尝试连接（如果 Studio 未运行会失败但不影响使用）
        studio_url = None
        if self.enable_studio:
            studio_url = "http://localhost:3000"
        
        init_agentscope(
            project="LingNexus",
            name="interactive_test",
            studio_url=studio_url,
            logging_path="./logs",
        )
    
    def _create_agent(self):
        """创建 Agent（使用渐进式披露方式）"""
        if self.agent is None:
            print("正在创建 Agent（渐进式披露模式）...")
            # 使用渐进式披露 Agent，推荐使用 qwen-max 作为 orchestrator
            # 渐进式披露模式统一使用 qwen-max，以获得最佳性能
            orchestrator_model = "qwen-max"
            self.agent = create_progressive_agent(
                model_name=orchestrator_model,
                temperature=0.3,  # orchestrator 使用较低温度
            )
            print(f"✅ Agent 创建成功（渐进式披露模式）")
            print(f"   Orchestrator: {orchestrator_model}")
            print(f"   💡 Agent 会自动按需加载 Skills 的完整指令，节省 tokens\n")
        return self.agent
    
    def _extract_response_text(self, response) -> str:
        """提取 Agent 响应的文本内容"""
        response_text = ""
        if hasattr(response, 'content'):
            if isinstance(response.content, list):
                for item in response.content:
                    if isinstance(item, dict) and item.get('type') == 'text':
                        response_text += item.get('text', '')
                    elif isinstance(item, str):
                        response_text += item
            else:
                response_text = str(response.content)
        else:
            response_text = str(response)
        return response_text
    
    async def _call_agent(self, user_input: str) -> str:
        """调用 Agent"""
        agent = self._create_agent()
        user_msg = Msg(name="user", role="user", content=user_input)
        response = await agent(user_msg)
        return self._extract_response_text(response)
    
    def _print_help(self):
        """打印帮助信息"""
        print("\n" + "=" * 60)
        print("命令帮助")
        print("=" * 60)
        print()
        print("命令:")
        print("  /help          - 显示此帮助信息")
        print("  /status        - 显示当前状态")
        print("  /mode <mode>   - 切换模式 (chat/test)")
        print("  /model <type>  - 切换模型 (qwen/deepseek)")
        print("  /execute <on/off> - 开启/关闭自动执行代码")
        print("  /studio <on/off>  - 开启/关闭 Studio")
        print("  /history       - 显示对话历史")
        print("  /clear         - 清空对话历史")
        print("  /files         - 列出生成的文件")
        print("  /view <file>   - 查看文件内容")
        print("  /exit          - 退出程序")
        print()
        print("使用说明:")
        print("  - 直接输入文本会发送给 Agent")
        print("  - 输入 / 开头的命令执行特殊操作")
        print("  - chat 模式: 普通对话，不执行代码")
        print("  - test 模式: 自动提取并执行代码")
        print("  - 当前使用渐进式披露模式（自动按需加载 Skills）")
        print()
        print("=" * 60)
        print()
    
    def _print_status(self):
        """打印当前状态"""
        print("\n" + "=" * 60)
        print("当前状态")
        print("=" * 60)
        print(f"  模式: {self.current_mode}")
        print(f"  模型: {self.model_type.value} ({self.model_name})")
        print(f"  自动执行代码: {'开启' if self.auto_execute_code else '关闭'}")
        print(f"  Studio: {'开启' if self.enable_studio else '关闭'}")
        print(f"  对话历史: {len(self.conversation_history)} 条")
        print("=" * 60)
        print()
    
    def _print_history(self):
        """打印对话历史"""
        if not self.conversation_history:
            print("暂无对话历史\n")
            return
        
        print("\n" + "=" * 60)
        print("对话历史")
        print("=" * 60)
        for i, (user, agent) in enumerate(self.conversation_history, 1):
            print(f"\n[{i}] 用户: {user[:50]}...")
            print(f"    Agent: {agent[:100]}...")
        print("=" * 60)
        print()
    
    def _handle_command(self, command: str) -> bool:
        """
        处理命令
        
        Returns:
            True 如果应该继续，False 如果应该退出
        """
        parts = command.strip().split()
        cmd = parts[0].lower()
        
        if cmd == "/help":
            self._print_help()
        
        elif cmd == "/status":
            self._print_status()
        
        elif cmd == "/mode":
            if len(parts) > 1:
                mode = parts[1].lower()
                if mode in ["chat", "test"]:
                    self.current_mode = mode
                    print(f"✅ 已切换到 {mode} 模式\n")
                else:
                    print("❌ 无效的模式，请使用 chat 或 test\n")
            else:
                print(f"当前模式: {self.current_mode}\n")
        
        elif cmd == "/model":
            if len(parts) > 1:
                model = parts[1].lower()
                if model in ["qwen", "deepseek"]:
                    self.model_type = ModelType.QWEN if model == "qwen" else ModelType.DEEPSEEK
                    self.model_name = "qwen-max" if model == "qwen" else "deepseek-chat"
                    self.agent = None  # 重置 Agent，下次调用时重新创建
                    print(f"✅ 已切换到 {model} 模型")
                    print(f"💡 注意: 渐进式披露模式统一使用 qwen-max 作为 orchestrator\n")
                else:
                    print("❌ 无效的模型，请使用 qwen 或 deepseek\n")
            else:
                print(f"当前模型: {self.model_type.value} ({self.model_name})")
                print(f"💡 注意: 渐进式披露模式统一使用 qwen-max 作为 orchestrator\n")
        
        elif cmd == "/execute":
            if len(parts) > 1:
                value = parts[1].lower()
                if value in ["on", "true", "1"]:
                    self.auto_execute_code = True
                    print("✅ 已开启自动执行代码\n")
                elif value in ["off", "false", "0"]:
                    self.auto_execute_code = False
                    print("✅ 已关闭自动执行代码\n")
                else:
                    print("❌ 无效的值，请使用 on/off\n")
            else:
                print(f"自动执行代码: {'开启' if self.auto_execute_code else '关闭'}\n")
        
        elif cmd == "/studio":
            if len(parts) > 1:
                value = parts[1].lower()
                if value in ["on", "true", "1"]:
                    self.enable_studio = True
                    self._init_agentscope()
                    print("✅ 已开启 Studio（需要重启程序生效）\n")
                elif value in ["off", "false", "0"]:
                    self.enable_studio = False
                    print("✅ 已关闭 Studio\n")
                else:
                    print("❌ 无效的值，请使用 on/off\n")
            else:
                print(f"Studio: {'开启' if self.enable_studio else '关闭'}\n")
        
        elif cmd == "/history":
            self._print_history()
        
        elif cmd == "/clear":
            self.conversation_history.clear()
            print("✅ 对话历史已清空\n")
        
        elif cmd == "/files":
            self._list_output_files()
        
        elif cmd == "/view":
            if len(parts) > 1:
                filename = parts[1]
                self._view_file_content(Path(filename))
            else:
                print("❌ 请指定文件名: /view <filename>\n")
        
        elif cmd == "/exit" or cmd == "/quit":
            return False
        
        else:
            print(f"❌ 未知命令: {cmd}，输入 /help 查看帮助\n")
        
        return True
    
    async def _process_user_input(self, user_input: str):
        """处理用户输入"""
        # 调用 Agent
        print("\n正在处理...")
        response_text = await self._call_agent(user_input)
        
        # 保存到历史
        self.conversation_history.append((user_input, response_text))
        
        # 显示响应
        print("\n" + "=" * 60)
        print("Agent 响应")
        print("=" * 60)
        print(response_text)
        print("=" * 60)
        
        # 根据模式处理
        if self.current_mode == "test" and self.auto_execute_code:
            # 检查是否有代码块
            has_code_block = response_text.count('```') >= 2

            if has_code_block:
                # 先检查有哪些语言的代码
                codes = extract_code_from_text(response_text)

                # 过滤掉看起来像是"执行命令"的 bash 代码
                # 如果 bash 代码只是展示执行命令（如 node -e, python 等），则跳过
                if 'bash' in codes:
                    bash_code = codes['bash']
                    # 如果 bash 代码只是单行命令（如 node -e, python 等），跳过
                    # 这些通常是 Agent 展示的执行结果，不需要再次执行
                    if bash_code and '\n' not in bash_code.strip():
                        # 检查是否是常见的代码执行命令
                        command_prefixes = ['node -e', 'python -c', 'python3 -c', 'php -r']
                        if any(bash_code.strip().startswith(prefix) for prefix in command_prefixes):
                            # 这是展示的执行命令，不是要执行的 bash 脚本
                            del codes['bash']

                if codes:
                    lang_names = list(codes.keys())
                    print("\n" + "=" * 60)
                    print(f"自动执行代码（检测到: {', '.join(lang_names)}）")
                    print("=" * 60)

                    # 使用多语言执行器（执行失败时保留临时文件用于调试）
                    result = await extract_and_execute_multi_language(response_text, keep_temp_file=True)

                    if result.get('code'):
                        lang = result.get('language', 'unknown')
                        print(f"✅ {lang.capitalize()} 代码提取成功")

                        if result['success']:
                            print(f"✅ {lang.capitalize()} 代码执行成功")
                            if result.get('output'):
                                print(f"输出:\n{result['output']}")
                            if result.get('returncode') is not None:
                                print(f"返回码: {result['returncode']}")
                        else:
                            print(f"❌ {lang.capitalize()} 代码执行失败")
                            if result.get('error'):
                                print(f"错误: {result['error']}")
                            if result.get('temp_file'):
                                print(f"💡 临时文件: {result['temp_file']}")
                            if result.get('returncode') is not None:
                                print(f"返回码: {result['returncode']}")
                else:
                    print("\n⚠️  检测到代码块，但无需执行的代码（可能是 Agent 展示的执行命令）")
                
                # 检查创建的文件
                current_dir = Path.cwd()
                docx_files = list(current_dir.glob("*.docx"))
                if docx_files:
                    print(f"\n✅ 发现 {len(docx_files)} 个 docx 文件:")
                    for f in docx_files[:5]:
                        size = f.stat().st_size
                        print(f"   - {f.name} ({size} 字节)")
        
        print()
    
    def _list_output_files(self):
        """列出输出文件"""
        print("\n" + "=" * 60)
        print("生成的文件")
        print("=" * 60)
        
        current_dir = Path.cwd()
        docx_files = list(current_dir.glob("*.docx"))
        
        if not docx_files:
            print("📁 当前目录暂无 docx 文件\n")
            return
        
        print(f"找到 {len(docx_files)} 个 docx 文件:\n")
        for i, file in enumerate(docx_files[:20], 1):  # 最多显示20个
            size = file.stat().st_size
            from datetime import datetime
            mtime = datetime.fromtimestamp(file.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")
            print(f"  {i}. {file.name}")
            print(f"     大小: {size} 字节")
            print(f"     时间: {mtime}")
            print()
        
        if len(docx_files) > 20:
            print(f"  ... 还有 {len(docx_files) - 20} 个文件未显示\n")
        
        print("💡 提示: 使用 /view <filename> 查看文件内容\n")
    
    def _view_file_content(self, file_path: Path):
        """查看文件内容"""
        if not file_path.exists():
            # 尝试在当前目录查找
            current_dir = Path.cwd()
            file_path = current_dir / file_path.name
        
        if not file_path.exists():
            print(f"❌ 文件不存在: {file_path}\n")
            return
        
        if not file_path.suffix == '.docx':
            print(f"❌ 只能查看 .docx 文件\n")
            return
        
        print("\n" + "=" * 60)
        print(f"文件内容: {file_path.name}")
        print("=" * 60)
        
        try:
            from docx import Document
            doc = Document(file_path)
            
            print(f"\n段落数量: {len(doc.paragraphs)}\n")
            print("内容:")
            print("-" * 60)
            
            for i, para in enumerate(doc.paragraphs[:20], 1):  # 最多显示20个段落
                if para.text.strip():
                    print(f"{i}. {para.text}")
            
            if len(doc.paragraphs) > 20:
                print(f"\n... 还有 {len(doc.paragraphs) - 20} 个段落未显示")
            
            print("-" * 60)
            print()
            
        except ImportError:
            print("⚠️  python-docx 未安装，无法读取文件内容")
            print("   安装命令: pip install python-docx\n")
        except Exception as e:
            print(f"❌ 读取文件时出错: {e}\n")
    
    async def run(self):
        """运行交互式测试"""
        # 初始化
        self._init_agentscope()
        
        # 打印欢迎信息
        print("\n" + "=" * 60)
        print("LingNexus 交互式测试工具")
        print("=" * 60)
        print()
        print("欢迎使用交互式测试工具！")
        print("✨ 当前使用渐进式披露模式（Progressive Disclosure）")
        print("   Agent 会自动按需加载 Skills，节省 tokens")
        print()
        print("输入 /help 查看帮助，输入 /exit 退出")
        print()
        self._print_status()
        
        # 主循环
        while True:
            try:
                # 获取用户输入
                prompt = f"[{self.current_mode}]> " if self.current_mode == "chat" else f"[{self.current_mode}+exec]> "
                user_input = input(prompt).strip()
                
                if not user_input:
                    continue
                
                # 处理命令
                if user_input.startswith("/"):
                    if not self._handle_command(user_input):
                        print("再见！\n")
                        break
                else:
                    # 处理普通输入
                    await self._process_user_input(user_input)
            
            except KeyboardInterrupt:
                print("\n\n中断操作，输入 /exit 退出\n")
            except EOFError:
                print("\n\n再见！\n")
                break
            except Exception as e:
                print(f"\n❌ 错误: {e}\n")
                import traceback
                traceback.print_exc()


async def main():
    """主函数（用于命令行调用）"""
    import argparse
    
    parser = argparse.ArgumentParser(description="LingNexus 交互式测试工具")
    parser.add_argument(
        '--model',
        choices=['qwen', 'deepseek'],
        default='qwen',
        help='模型类型 (默认: qwen)'
    )
    parser.add_argument(
        '--model-name',
        type=str,
        default=None,
        help='模型名称（如 qwen-max, deepseek-chat）'
    )
    parser.add_argument(
        '--mode',
        choices=['chat', 'test'],
        default='test',
        help='初始模式 (默认: test)'
    )
    parser.add_argument(
        '--no-execute',
        action='store_true',
        help='不自动执行代码'
    )
    parser.add_argument(
        '--studio',
        action='store_true',
        help='启用 Studio'
    )
    parser.add_argument(
        '--no-studio',
        action='store_true',
        help='禁用 Studio（覆盖环境变量）'
    )
    
    args = parser.parse_args()
    
    # 默认开启 Studio（如果环境变量设置了 ENABLE_STUDIO=true）
    # 或者通过 --studio 参数显式开启
    # 可以通过 --no-studio 显式禁用
    default_studio = os.getenv("ENABLE_STUDIO", "false").lower() == "true"
    enable_studio = args.studio or (default_studio and not args.no_studio)
    
    model_type = ModelType.QWEN if args.model == 'qwen' else ModelType.DEEPSEEK
    model_name = args.model_name or ("qwen-max" if args.model == 'qwen' else "deepseek-chat")
    
    tester = InteractiveTester(
        model_type=model_type,
        model_name=model_name,
        auto_execute_code=not args.no_execute,
        enable_studio=enable_studio,
    )
    tester.current_mode = args.mode
    
    await tester.run()


if __name__ == "__main__":
    asyncio.run(main())

